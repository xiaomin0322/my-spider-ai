"""
知源中医网 - 解表药 > 发散风寒药（白芷、防风）信息抓取 + Word文档生成

功能：
  从知源中医网抓取"解表药"分类概述，以及"发散风寒药"下白芷、防风两味药的详细信息，
  生成带层级导航的Word文档。

数据来源：
  https://www.zhiyuanzhongyi.com/traditional（知源中医网 - 常用中药）

输出文件：
  C:\Users\zengmin.zhang\AI\知源中医_解表药_发散风寒药.docx

文档结构（3级导航）：
  H1  一、解表药
    H2  （一）发散风寒药
      H3  分类概述（定义、功用、分类、配伍、用药注意）
      H3  1. 白芷（基本信息、临床应用、相关配伍、相关方剂、炮制）
      H3  2. 防风（基本信息、临床应用、相关配伍、相关方剂、炮制）

扩展方式：
  修改 HERBS 列表，添加其他药物的 id 和 url 即可抓取更多药物。
  药物 id 可在 https://www.zhiyuanzhongyi.com/traditional 页面的链接中找到。
"""

import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ===== 配置 =====
OUTPUT_DIR = r'C:\Users\zengmin.zhang\AI'
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, '知源中医_解表药_发散风寒药.docx')
BASE_URL = 'https://www.zhiyuanzhongyi.com'

# 要抓取的药物列表（id 从网站链接中获取，如 /traditionaldetails?id=41）
HERBS = [
    {'name': '白芷', 'id': 41, 'url': '/traditionaldetails?id=41'},
    {'name': '防风', 'id': 61, 'url': '/traditionaldetails?id=61'},
]

# 分类页面URL（解表药的分类详情页）
CATEGORY_URL = '/traditionaltypedetails?id=1'

HEADERS = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}


# ============================================================
#  第一部分：网页抓取
# ============================================================

def fetch_page(url):
    """请求网页并返回 BeautifulSoup 对象"""
    resp = requests.get(url, timeout=15, headers=HEADERS)
    resp.encoding = 'utf-8'
    return BeautifulSoup(resp.text, 'lxml')


def parse_category_page(soup):
    """
    解析分类页面（如"解表药"），提取分类描述信息。

    返回 dict，包含：
      name          — 分类名称（如"解表药"）
      definition    — 定义
      functions     — 功用
      classification — 分类说明
      compatibility — 配伍原则
      cautions      — 用药注意
    """
    text = soup.get_text(separator='\n', strip=True)
    lines = text.split('\n')

    info = {
        'name': '',
        'definition': '',
        'functions': '',
        'classification': '',
        'compatibility': '',
        'cautions': '',
    }

    # 提取分类名称
    for i, line in enumerate(lines):
        if line == '解表药' and i < 5:
            info['name'] = '解表药'
            break

    # 按"定义/功用/分类/配伍/用药注意"分段提取
    current_section = None
    section_text = []

    for line in lines:
        if line.startswith('定义:') or line.startswith('定义：'):
            if current_section:
                info[current_section] = '\n'.join(section_text).strip()
            current_section = 'definition'
            section_text = [line.split(':', 1)[-1].split('：', 1)[-1].strip()]
        elif line.startswith('功用:') or line.startswith('功用：'):
            if current_section:
                info[current_section] = '\n'.join(section_text).strip()
            current_section = 'functions'
            section_text = [line.split(':', 1)[-1].split('：', 1)[-1].strip()]
        elif line.startswith('分类:') or line.startswith('分类：'):
            if current_section:
                info[current_section] = '\n'.join(section_text).strip()
            current_section = 'classification'
            section_text = [line.split(':', 1)[-1].split('：', 1)[-1].strip()]
        elif line.startswith('配伍:') or line.startswith('配伍：'):
            if current_section:
                info[current_section] = '\n'.join(section_text).strip()
            current_section = 'compatibility'
            section_text = [line.split(':', 1)[-1].split('：', 1)[-1].strip()]
        elif line.startswith('用药注意:') or line.startswith('用药注意：'):
            if current_section:
                info[current_section] = '\n'.join(section_text).strip()
            current_section = 'cautions'
            section_text = [line.split(':', 1)[-1].split('：', 1)[-1].strip()]
        elif current_section and line and line not in ('常用解表药:', '常用解表药：'):
            section_text.append(line)

    if current_section:
        info[current_section] = '\n'.join(section_text).strip()

    return info


def parse_herb_page(soup):
    """
    解析药物详情页面，提取结构化信息。

    返回 dict，包含：
      基本信息：name, pinyin, source, aliases, nature_taste_channel, effects, introduction, usage_dosage, cautions
      扩展信息：applications（临床应用）, pairings（相关配伍）, prescriptions（相关方剂）, processing_*（炮制）

    页面结构：
      药品名称 → 始载于 → 别名 → 性味归经 → 功效 → 药材简介 → 用法用量 → 注意事项
      → 应用（编号列表）→ 相关配伍（配对说明）→ 方剂（方名+组成+功能）→ 炮制
    """
    text = soup.get_text(separator='\n', strip=True)
    lines = text.split('\n')

    herb = {
        'name': '', 'pinyin': '', 'source': '', 'aliases': '',
        'nature_taste_channel': '', 'effects': '', 'introduction': '',
        'usage_dosage': '', 'cautions': '',
        'applications': [], 'pairings': [], 'prescriptions': [],
        'processing_modern': '', 'processing_ancient': '',
    }

    # ---- 提取基本信息字段 ----
    # 页面格式：字段名 → "：" → 字段值（占3行）
    for i, line in enumerate(lines):
        if line == '药品名称' and i + 2 < len(lines):
            val = lines[i + 2]
            herb['name'] = val.split('[')[0].strip() if '[' in val else val
            herb['pinyin'] = val
        elif line == '始载于' and i + 2 < len(lines):
            herb['source'] = lines[i + 2]
        elif line == '别名' and i + 2 < len(lines):
            herb['aliases'] = lines[i + 2]
        elif line == '性味归经' and i + 2 < len(lines):
            herb['nature_taste_channel'] = lines[i + 2]
        elif line == '功效' and i + 2 < len(lines) and lines[i + 1] == '：':
            herb['effects'] = lines[i + 2]
        elif line == '药材简介' and i + 2 < len(lines):
            herb['introduction'] = lines[i + 2]
        elif line == '用法用量' and i + 2 < len(lines):
            # 用法用量可能有多行，读到下一个字段为止
            usage = []
            j = i + 2
            while j < len(lines) and lines[j] not in ('注意事项', '应用', '炮制', '基本信息', '药品名称'):
                if lines[j].strip():
                    usage.append(lines[j])
                j += 1
            herb['usage_dosage'] = '\n'.join(usage)
        elif line == '注意事项' and i + 2 < len(lines):
            cautions = []
            j = i + 2
            while j < len(lines) and lines[j] not in ('应用', '炮制', '基本信息', '药品名称'):
                if lines[j].strip():
                    cautions.append(lines[j])
                j += 1
            herb['cautions'] = '\n'.join(cautions)

    # ---- 提取临床应用（编号列表：1. xxx  2. xxx）----
    in_app = False
    for line in lines:
        if line == '应用':
            in_app = True
            continue
        if line in ('相关配伍', '炮制'):
            in_app = False
            continue
        if in_app and line and line[0].isdigit() and '.' in line[:3]:
            herb['applications'].append(line)

    # ---- 提取相关配伍（如"白芷配细辛"）----
    in_pair = False
    current_pair = None
    for line in lines:
        if line == '相关配伍':
            in_pair = True
            continue
        if line == '炮制':
            in_pair = False
            current_pair = None
            continue
        if in_pair:
            # 配伍标题：含"配"字，较短，不含方剂名
            if '配' in line and len(line) < 30 and '汤' not in line and '散' not in line and '丸' not in line:
                current_pair = {'title': line, 'content': ''}
                herb['pairings'].append(current_pair)
            elif current_pair and line not in ('相关配伍', '炮制'):
                # 遇到方剂名则结束当前配伍
                is_rx = ('汤' in line or '散' in line or '丸' in line or '饮' in line) and len(line) < 20
                if is_rx:
                    current_pair = None
                elif current_pair:
                    if current_pair['content']:
                        current_pair['content'] += '\n' + line
                    else:
                        current_pair['content'] = line

    # ---- 提取相关方剂（方名 + 药物组成 + 功能与主治）----
    seen_rx = set()
    for i, line in enumerate(lines):
        if ('汤' in line or '散' in line or '丸' in line or '饮' in line) and \
           '药物组成' not in line and '功能与主治' not in line and \
           len(line) < 20 and line not in seen_rx:
            j = i + 1
            comp = ''
            func = ''
            while j < len(lines) and j < i + 5:
                if '药物组成' in lines[j]:
                    comp = lines[j].split(':', 1)[-1].split('：', 1)[-1].strip()
                elif '功能与主治' in lines[j]:
                    func_parts = [lines[j].split(':', 1)[-1].split('：', 1)[-1].strip()]
                    # 功能描述可能跨多行
                    k = j + 1
                    while k < len(lines):
                        nl = lines[k].strip()
                        if not nl or '药物组成' in nl or '功能与主治' in nl:
                            break
                        if nl in ('相关配伍', '炮制', '现代炮制', '古法炮制'):
                            break
                        if ('汤' in nl or '散' in nl or '丸' in nl or '饮' in nl) and len(nl) < 20:
                            break
                        func_parts.append(nl)
                        k += 1
                    func = ''.join(func_parts)
                    break
                j += 1
            if comp or func:
                herb['prescriptions'].append({'name': line, 'composition': comp, 'function': func})
                seen_rx.add(line)

    # ---- 提取炮制方法（现代炮制 + 古法炮制）----
    if '炮制' in text:
        proc_section = text.split('炮制')[1]
        if '扫码下载' in proc_section:
            proc_section = proc_section.split('扫码下载')[0]
        if '现代炮制' in proc_section:
            parts = proc_section.split('现代炮制')
            modern = parts[1] if len(parts) > 1 else ''
            if '古法炮制' in modern:
                modern = modern.split('古法炮制')[0]
            herb['processing_modern'] = modern.strip().lstrip('：:').strip()
        if '古法炮制' in proc_section:
            ancient = proc_section.split('古法炮制')[1]
            herb['processing_ancient'] = ancient.strip().lstrip('：:').strip()

    return herb


# ============================================================
#  第二部分：Word文档生成
# ============================================================

def set_cell_shading(cell, color):
    """设置Word表格单元格的背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading_elm.append(shading)


def add_heading(doc, text, level=1):
    """添加带统一颜色的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)  # 深绿色
    return h


def add_section_title(doc, text):
    """添加带下划线的小节标题（如"临床应用"、"相关配伍"）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)  # 蓝色
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    # 底部蓝色细线
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '0d47a1'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(category_info, herbs_data):
    """
    生成Word文档。

    参数：
      category_info — 分类概述信息（来自 parse_category_page）
      herbs_data    — 药物详情列表（来自 parse_herb_page）

    文档结构：
      H1  一、解表药
        H2  （一）发散风寒药
          H3  分类概述
          H3  1. 白芷
          H3  2. 防风
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    # ===== 文档标题 =====
    title = doc.add_heading('知源中医 - 常用中药', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)

    doc.add_paragraph()

    # ===== H1: 大分类 =====
    add_heading(doc, '一、解表药', level=1)

    # ===== H2: 子分类 =====
    add_heading(doc, '（一）发散风寒药', level=2)

    # ===== H3: 分类概述 =====
    add_heading(doc, '分类概述', level=3)

    if category_info.get('definition'):
        add_section_title(doc, '定义')
        p = doc.add_paragraph(category_info['definition'])
        p.paragraph_format.first_line_indent = Cm(0.75)

    if category_info.get('functions'):
        add_section_title(doc, '功用')
        p = doc.add_paragraph(category_info['functions'])
        p.paragraph_format.first_line_indent = Cm(0.75)

    if category_info.get('classification'):
        add_section_title(doc, '分类')
        p = doc.add_paragraph(category_info['classification'])
        p.paragraph_format.first_line_indent = Cm(0.75)

    if category_info.get('compatibility'):
        add_section_title(doc, '配伍')
        p = doc.add_paragraph(category_info['compatibility'])
        p.paragraph_format.first_line_indent = Cm(0.75)

    if category_info.get('cautions'):
        add_section_title(doc, '用药注意')
        p = doc.add_paragraph(category_info['cautions'])
        p.paragraph_format.first_line_indent = Cm(0.75)

    doc.add_paragraph()

    # ===== H3: 各药物详情 =====
    for idx, herb in enumerate(herbs_data, 1):
        add_heading(doc, f'{idx}. {herb["name"]}', level=3)

        # ---- 基本信息表格 ----
        add_section_title(doc, '基本信息')
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(3.5)
        table.columns[1].width = Cm(13)

        fields = [
            ('药品名称', herb.get('name', '') + (' ' + herb['pinyin'] if herb.get('pinyin') else '')),
            ('始载于', herb.get('source', '')),
            ('别名', herb.get('aliases', '')),
            ('性味归经', herb.get('nature_taste_channel', '')),
            ('功效', herb.get('effects', '')),
            ('药材简介', herb.get('introduction', '')),
            ('用法用量', herb.get('usage_dosage', '')),
            ('注意事项', herb.get('cautions', '')),
        ]

        for label, value in fields:
            if value:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = value
                # 标签列加粗 + 绿色背景
                for p in row.cells[0].paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                set_cell_shading(row.cells[0], 'e8f5e9')
                # 值列正常字体
                for p in row.cells[1].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph()

        # ---- 临床应用 ----
        if herb.get('applications'):
            add_section_title(doc, '临床应用')
            for app in herb['applications']:
                p = doc.add_paragraph(app, style='List Number')
                p.paragraph_format.space_after = Pt(4)

        # ---- 相关配伍 ----
        if herb.get('pairings'):
            add_section_title(doc, '相关配伍')
            for pairing in herb['pairings']:
                p = doc.add_paragraph()
                run = p.add_run(pairing['title'])
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x2e, 0x7d, 0x32)  # 绿色
                if pairing.get('content'):
                    p2 = doc.add_paragraph(pairing['content'])
                    p2.paragraph_format.first_line_indent = Cm(0.75)
                    p2.paragraph_format.space_after = Pt(6)

        # ---- 相关方剂 ----
        if herb.get('prescriptions'):
            add_section_title(doc, '相关方剂')
            for rx in herb['prescriptions']:
                # 方剂名（橙色加粗）
                p = doc.add_paragraph()
                run = p.add_run(rx['name'])
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xbf, 0x36, 0x0c)
                # 药物组成
                if rx.get('composition'):
                    p2 = doc.add_paragraph()
                    run_label = p2.add_run('药物组成：')
                    run_label.bold = True
                    run_label.font.size = Pt(10)
                    run_val = p2.add_run(rx['composition'])
                    run_val.font.size = Pt(10)
                    p2.paragraph_format.space_after = Pt(2)
                # 功能与主治
                if rx.get('function'):
                    p3 = doc.add_paragraph()
                    run_label = p3.add_run('功能与主治：')
                    run_label.bold = True
                    run_label.font.size = Pt(10)
                    run_val = p3.add_run(rx['function'])
                    run_val.font.size = Pt(10)
                    p3.paragraph_format.space_after = Pt(8)

        # ---- 炮制 ----
        if herb.get('processing_modern') or herb.get('processing_ancient'):
            add_section_title(doc, '炮制')
            if herb.get('processing_modern'):
                p = doc.add_paragraph()
                run = p.add_run('现代炮制：')
                run.bold = True
                run.font.size = Pt(10.5)
                run2 = p.add_run(herb['processing_modern'])
                run2.font.size = Pt(10)
            if herb.get('processing_ancient'):
                p = doc.add_paragraph()
                run = p.add_run('古法炮制：')
                run.bold = True
                run.font.size = Pt(10.5)
                run2 = p.add_run(herb['processing_ancient'])
                run2.font.size = Pt(10)

        if idx < len(herbs_data):
            doc.add_page_break()

    # ===== 页脚 =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('数据来源：知源中医网 (www.zhiyuanzhongyi.com)')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.save(OUTPUT_DOCX)
    print(f'文档已生成: {OUTPUT_DOCX}')
    print(f'文件大小: {os.path.getsize(OUTPUT_DOCX) / 1024:.1f} KB')


# ============================================================
#  第三部分：主函数
# ============================================================

def main():
    """
    主流程：抓取分类页面 → 抓取药物详情 → 生成Word文档
    """
    print("=" * 50)
    print("知源中医 - 解表药 > 发散风寒药 信息抓取")
    print("=" * 50)

    # 1. 抓取分类概述页面
    print("\n1. 抓取分类页面（解表药）...")
    cat_soup = fetch_page(BASE_URL + CATEGORY_URL)
    category_info = parse_category_page(cat_soup)
    print(f"   分类: {category_info['name']}")
    print(f"   定义: {category_info['definition'][:60]}...")

    # 2. 逐个抓取药物详情
    herbs_data = []
    for herb_item in HERBS:
        print(f"\n2. 抓取药物详情：{herb_item['name']}...")
        herb_soup = fetch_page(BASE_URL + herb_item['url'])
        herb_data = parse_herb_page(herb_soup)
        herb_data['name'] = herb_item['name']
        herbs_data.append(herb_data)
        print(f"   功效: {herb_data['effects'][:60]}...")
        print(f"   临床应用: {len(herb_data['applications'])} 条")
        print(f"   相关配伍: {len(herb_data['pairings'])} 条")
        print(f"   相关方剂: {len(herb_data['prescriptions'])} 条")

    # 3. 生成Word文档
    print(f"\n3. 生成Word文档...")
    generate_docx(category_info, herbs_data)

    print("\n完成！")


if __name__ == '__main__':
    main()
