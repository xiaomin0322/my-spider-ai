"""
民间中医网 · 经方方药版块 · 全部博文 → Word文档

功能：
  读取抓取的帖子数据（thread_results.json），生成带颜色标注的Word文档。
  每篇博文包含：主贴内容 + 楼主回复，对话格式用颜色区分发言者。

数据来源：
  http://www.ngotcmszh.com/forum-180-1.html（经方方药版块，共1424篇）

输入文件：
  %TEMP%/thread_results.json  — 由抓取脚本生成的帖子数据

输出文件：
  C:\Users\zengmin.zhang\民间中医网经方方药全部博文.docx

颜色规则：
  红色 = 水中火（版主/老师）
  蓝色 = 患者及家属
  棕色 = 其他发言者
  灰色 = 日期、时间戳
"""

import os
import sys
import json

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 文件路径 =====
TEMP_DIR = os.environ.get('TEMP', os.environ.get('TMP', '/tmp'))
RESULTS_JSON = os.path.join(TEMP_DIR, 'thread_results.json')
OUTPUT_PATH = r'C:\Users\zengmin.zhang\民间中医网经方方药全部博文.docx'

# ===== 发言者颜色配置 =====
SPEAKER_COLORS = {
    '水中火': RGBColor(0xC0, 0x39, 0x2B),    # 红色 - 版主
    '患者': RGBColor(0x27, 0x63, 0xA3),        # 蓝色
    '患者妻子': RGBColor(0x27, 0x63, 0xA3),
    '患者媳妇': RGBColor(0x27, 0x63, 0xA3),
    '患者家属': RGBColor(0x27, 0x63, 0xA3),
}
DEFAULT_SPEAKER_COLOR = RGBColor(0x6C, 0x3C, 0x0A)  # 棕色 - 其他
DATE_COLOR = RGBColor(0x80, 0x80, 0x80)               # 灰色 - 日期


def add_conversation_to_doc(doc, conversations):
    """
    将对话内容添加到文档。

    conversations 是一个列表，每个元素是 dict，包含以下 type：
      - 'date'    : 日期分隔行，如 "2024-01-15"
      - 'speaker' : 发言者+时间，如 {name: '水中火', time: '14:30', message: '...'}
      - 'message' : 普通消息文本
      - 'text'    : 纯文本段落（非对话格式）
    """
    current_para = None

    for item in conversations:
        # 日期行：居中显示灰色分隔线
        if item['type'] == 'date':
            para = doc.add_paragraph()
            date_run = para.add_run(f"───── {item['content']} ─────")
            date_run.font.color.rgb = DATE_COLOR
            date_run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_para = None

        # 发言者行：带颜色的名字 + 时间 + 消息内容
        elif item['type'] == 'speaker':
            para = doc.add_paragraph()
            speaker_name = item['name']
            speaker_time = item.get('time', '')

            # 发言者名字（带颜色加粗）
            color = SPEAKER_COLORS.get(speaker_name, DEFAULT_SPEAKER_COLOR)
            name_run = para.add_run(f"【{speaker_name}】")
            name_run.bold = True
            name_run.font.color.rgb = color
            name_run.font.size = Pt(11)

            # 时间戳（灰色小字）
            if speaker_time:
                time_run = para.add_run(f" {speaker_time}")
                time_run.font.color.rgb = DATE_COLOR
                time_run.font.size = Pt(9)

            # 消息内容（紧跟发言者后面）
            if 'message' in item:
                msg_run = para.add_run(f"\n{item['message']}")
                msg_run.font.size = Pt(11)

            current_para = para

        # 普通消息：追加到当前发言者段落，或新建独立段落
        elif item['type'] == 'message':
            if current_para:
                msg_run = current_para.add_run(f"\n{item['content']}")
                msg_run.font.size = Pt(11)
            else:
                para = doc.add_paragraph(item['content'])
                para.paragraph_format.first_line_indent = Pt(22)
            current_para = None

        # 纯文本：普通段落，首行缩进
        elif item['type'] == 'text':
            if item['content']:
                para = doc.add_paragraph(item['content'])
                para.paragraph_format.first_line_indent = Pt(22)
            current_para = None


def create_doc(results, output_path):
    """
    根据帖子数据生成Word文档。

    参数：
      results     — 帖子数据列表，每个元素包含 title/url/main_author/main_post/author_replies
      output_path — 输出文件路径

    文档结构：
      标题 → 统计信息 → 颜色说明 → 每篇博文（标题+主贴+楼主回帖）
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)

    # ===== 文档标题 =====
    title_para = doc.add_heading('民间中医网 · 经方方药版块 · 全部博文', level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 统计信息 =====
    total_replies = sum(len(r.get('author_replies', [])) for r in results)
    info_para = doc.add_paragraph()
    info_para.add_run(f"共收录 {len(results)} 篇博文，{total_replies} 条楼主回复").font.color.rgb = DATE_COLOR
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 颜色说明 =====
    legend_para = doc.add_paragraph()
    legend_para.add_run('颜色说明：').bold = True
    for name, color in SPEAKER_COLORS.items():
        legend_para.add_run(f'  ■ {name}').font.color.rgb = color
    legend_para.add_run('  ■ 其他').font.color.rgb = DEFAULT_SPEAKER_COLOR

    # ===== 逐篇博文 =====
    for idx, thread in enumerate(results, 1):
        if idx > 1:
            doc.add_page_break()  # 每篇博文从新页开始

        title = thread.get('title', '未知标题')
        url = thread.get('url', '')
        main_author = thread.get('main_author', '')
        main_post = thread.get('main_post', [])
        author_replies = thread.get('author_replies', [])

        # 帖子标题
        doc.add_heading(f'{idx}. {title}', level=1)

        # 信息行：链接 + 作者
        info_para = doc.add_paragraph()
        info_para.add_run('链接：').bold = True
        info_para.add_run(url)
        info_para.add_run('    作者：').bold = True
        author_run = info_para.add_run(main_author)
        author_run.font.color.rgb = SPEAKER_COLORS.get(main_author, DEFAULT_SPEAKER_COLOR)
        author_run.bold = True

        # ===== 主贴内容 =====
        doc.add_heading('【主贴】', level=2)

        for post in main_post:
            # 作者名 + 发帖时间
            post_info = doc.add_paragraph()
            author_name_run = post_info.add_run(f"[{post['author']}]")
            author_name_run.bold = True
            author_name_run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

            if post.get('time'):
                time_run = post_info.add_run(f"  {post['time']}")
                time_run.font.color.rgb = DATE_COLOR
                time_run.font.size = Pt(9)

            # 内容：如果是对话格式则用颜色标注，否则普通文本
            if post.get('is_conversation') and isinstance(post['content'], list):
                add_conversation_to_doc(doc, post['content'])
            else:
                content_para = doc.add_paragraph(post['content'])
                content_para.paragraph_format.first_line_indent = Pt(22)

        # ===== 楼主回帖 =====
        if author_replies:
            doc.add_heading('【楼主回帖】', level=2)

            for post in author_replies:
                # 楼层号 + 作者名 + 时间
                post_info = doc.add_paragraph()
                floor_run = post_info.add_run(f"{post['floor']} ")
                floor_run.bold = True
                floor_run.font.size = Pt(10)

                author_name_run = post_info.add_run(f"[{post['author']}]")
                author_name_run.bold = True
                author_name_run.font.color.rgb = SPEAKER_COLORS.get(post['author'], DEFAULT_SPEAKER_COLOR)

                if post.get('time'):
                    time_run = post_info.add_run(f"  {post['time']}")
                    time_run.font.color.rgb = DATE_COLOR
                    time_run.font.size = Pt(9)

                # 内容
                if post.get('is_conversation') and isinstance(post['content'], list):
                    add_conversation_to_doc(doc, post['content'])
                else:
                    content_para = doc.add_paragraph(post['content'])
                    content_para.paragraph_format.first_line_indent = Pt(22)

                # 分隔线（区分不同回帖）
                sep_para = doc.add_paragraph()
                sep_para.add_run('─' * 50).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.save(output_path)
    return output_path


def main():
    """主函数：加载数据 → 生成文档"""
    print("=" * 60)
    print("生成Word文档")
    print("=" * 60)

    # 检查数据文件是否存在
    if not os.path.exists(RESULTS_JSON):
        print(f"错误: 找不到结果文件 {RESULTS_JSON}")
        print("请先运行抓取脚本生成数据")
        return

    # 加载帖子数据
    with open(RESULTS_JSON, 'r', encoding='utf-8') as f:
        results = json.load(f)

    print(f"加载 {len(results)} 个帖子数据")

    # 生成文档
    output_path = create_doc(results, OUTPUT_PATH)

    file_size = os.path.getsize(output_path)
    print(f"\n文档已生成: {output_path}")
    print(f"文件大小: {file_size / 1024 / 1024:.1f} MB")


if __name__ == '__main__':
    main()
