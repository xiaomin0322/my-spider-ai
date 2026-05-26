"""
知源中医网 - 全量中药抓取 + Word文档生成

功能：
  从 https://www.zhiyuanzhongyi.com/traditional 抓取全部439味中药的详细信息，
  按照网站的21个大分类（解表药、清热药、泻下药...）和子分类组织，
  生成带3级导航的Word文档。

主要流程：
  1. 访问主页，获取所有药物的URL映射（药名→链接）
  2. 逐个抓取每味药的详情页，解析为结构化数据
  3. 按 分类→子分类→药物 的层级生成Word文档

输入：
  无外部文件依赖，直接从网站抓取

输出：
  - C:\Users\zengmin.zhang\AI\知源中医_常用中药大全.docx
  - %TEMP%/herbs_progress.json（断点续传用的进度文件）

特性：
  - 断点续传：进度保存在 herbs_progress.json，重新运行会跳过已抓取的药物
  - 限速控制：每药间隔3秒，避免被封IP
  - 每味药包含：基本信息、临床应用、相关配伍、相关方剂、炮制方法
"""

import os
import sys
import re
import json
import time
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

# ===== 文件路径配置 =====
OUTPUT_DIR = r'C:\Users\zengmin.zhang\AI'
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, '知源中医_常用中药大全.docx')  # 最终输出的Word文档
TEMP_DIR = os.environ.get('TEMP', '/tmp')
PROGRESS_JSON = os.path.join(TEMP_DIR, 'herbs_progress.json')  # 抓取进度（断点续传用）

# ===== 网络请求配置 =====
BASE_URL = 'https://www.zhiyuanzhongyi.com'
HEADERS = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
REQUEST_DELAY = 3  # 每次请求间隔（秒），避免被封IP

# ===== 中药分类结构 =====
# 按照网站页面的实际排列顺序组织
# 结构：大分类(id/name) → 子分类(name) → 药物列表(herbs)
# 药物名称必须与网站上的名称完全一致（用于匹配URL和数据）
CATEGORIES = [
    {
        'id': 1, 'name': '解表药',
        'subcategories': [
            {'name': '发散风寒药', 'herbs': ['白芷', '防风', '羌活', '细辛', '生姜', '藁本', '桂枝', '紫苏叶', '辛夷', '谷精草', '苍耳子', '荆芥', '香薷', '麻黄', '葱白', '西河柳', '胡荽']},
            {'name': '发散风热药', 'herbs': ['升麻', '柴胡', '葛根', '桑叶', '菊花', '牛蒡子', '蔓荆子', '蝉蜕', '木贼', '薄荷', '浮萍']},
        ]
    },
    {
        'id': 2, 'name': '清热药',
        'subcategories': [
            {'name': '清热泻火药', 'herbs': ['天花粉', '知母', '芦根', '夏枯草', '石膏', '决明子', '青葙子', '栀子', '淡竹叶', '寒水石', '密蒙花', '鸭跖草']},
            {'name': '清热燥湿药', 'herbs': ['苦参', '黄芩', '黄连', '龙胆', '白鲜皮', '秦皮', '黄柏']},
            {'name': '清热解毒药', 'herbs': ['土茯苓', '山豆根', '白头翁', '白蔹', '板蓝根', '金荞麦', '重楼', '射干', '漏芦', '贯众', '大血藤', '大青叶', '金银花', '野菊花', '木蝴蝶', '青果', '鸦胆子', '青黛', '马勃', '半边莲', '白花蛇舌草', '穿心莲', '鱼腥草', '蒲公英', '千里光', '四季青', '地锦草', '马齿苋', '连翘', '拳参', '败酱草', '熊胆粉', '绿豆', '紫花地丁', '山慈菇']},
            {'name': '清热凉血药', 'herbs': ['玄参', '赤芍', '紫草', '牡丹皮', '水牛角', '生地黄']},
            {'name': '清虚热药', 'herbs': ['白薇', '胡黄连', '银柴胡', '地骨皮', '青蒿']},
        ]
    },
    {
        'id': 3, 'name': '泻下药',
        'subcategories': [
            {'name': '攻下药', 'herbs': ['大黄', '番泻叶', '芒硝', '芦荟']},
            {'name': '润下药', 'herbs': ['火麻仁', '郁李仁', '松子仁']},
            {'name': '峻下逐水药', 'herbs': ['甘遂', '京大戟', '商陆', '芫花', '千金子', '巴豆霜', '牵牛子']},
        ]
    },
    {
        'id': 4, 'name': '祛风湿药',
        'subcategories': [
            {'name': '祛风寒湿药', 'herbs': ['川乌', '威灵仙', '徐长卿', '独活', '青风藤', '木瓜', '路路通', '乌梢蛇', '蕲蛇', '伸筋草', '丁公藤', '海风藤', '蚕沙', '油松节', '昆明山海棠']},
            {'name': '祛风湿热药', 'herbs': ['防己', '秦艽', '桑枝', '络石藤', '丝瓜络', '老鹳草', '豨莶草', '穿山龙', '臭梧桐', '海桐皮', '雷公藤']},
            {'name': '祛风湿强筋骨药', 'herbs': ['千年健', '狗脊', '五加皮', '桑寄生', '雪莲花']},
        ]
    },
    {
        'id': 5, 'name': '化湿药',
        'subcategories': [
            {'name': '化湿药', 'herbs': ['苍术', '厚朴', '豆蔻', '砂仁', '佩兰', '广藿香', '草豆蔻', '草果']},
        ]
    },
    {
        'id': 6, 'name': '利水渗湿药',
        'subcategories': [
            {'name': '利水消肿药', 'herbs': ['泽泻', '香加皮', '薏苡仁', '茯苓', '猪苓', '冬瓜皮', '枳椇子', '玉米须', '葫芦']},
            {'name': '利尿通淋药', 'herbs': ['木通', '通草', '石韦', '滑石', '地肤子', '车前子', '海金沙', '灯心草', '瞿麦', '冬葵子', '萹蓄', '萆薢']},
            {'name': '利湿退黄药', 'herbs': ['虎杖', '垂盆草', '金钱草', '茵陈', '鸡骨草', '珍珠草', '地耳草']},
        ]
    },
    {
        'id': 7, 'name': '温里药',
        'subcategories': [
            {'name': '温里药', 'herbs': ['附子', '高良姜', '干姜', '肉桂', '丁香', '小茴香', '吴茱萸', '花椒', '胡椒', '荜茇', '荜澄茄', '荜苃']},
        ]
    },
    {
        'id': 8, 'name': '理气药',
        'subcategories': [
            {'name': '理气药', 'herbs': ['木香', '甘松', '香附', '乌药', '薤白', '沉香', '檀香', '玫瑰花', '刀豆', '大腹皮', '川楝子', '佛手', '青皮', '枳实', '柿蒂', '九香虫', '荔枝核', '娑罗子', '陈皮', '香橼', '梅花']},
        ]
    },
    {
        'id': 9, 'name': '消食药',
        'subcategories': [
            {'name': '消食药', 'herbs': ['山楂', '莱菔子', '麦芽', '稻芽', '鸡内金', '六神曲']},
        ]
    },
    {
        'id': 10, 'name': '驱虫药',
        'subcategories': [
            {'name': '驱虫药', 'herbs': ['苦楝皮', '使君子', '榧子', '槟榔', '鹤虱', '雷丸', '芜荑', '南瓜子', '鹤草芽']},
        ]
    },
    {
        'id': 11, 'name': '止血药',
        'subcategories': [
            {'name': '凉血止血药', 'herbs': ['白茅根', '地榆', '侧柏叶', '槐花', '大蓟', '小蓟', '羊蹄', '苎麻根']},
            {'name': '化瘀止血药', 'herbs': ['三七', '茜草', '蒲黄', '花蕊石']},
            {'name': '收敛止血药', 'herbs': ['白及', '仙鹤草', '紫珠叶', '血余炭', '藕节']},
            {'name': '温经止血药', 'herbs': ['艾叶', '炮姜', '灶心土']},
        ]
    },
    {
        'id': 12, 'name': '活血化瘀药',
        'subcategories': [
            {'name': '活血止痛药', 'herbs': ['川芎', '延胡索', '姜黄', '郁金', '降香', '没药', '乳香', '五灵脂']},
            {'name': '活血调经药', 'herbs': ['丹参', '牛膝', '鸡血藤', '月季花', '红花', '凌霄花', '王不留行', '益母草', '泽兰', '桃仁']},
            {'name': '活血疗伤药', 'herbs': ['骨碎补', '苏木', '自然铜', '血竭', '儿茶', '土鳖虫', '刘寄奴', '马钱子']},
            {'name': '破血消癥药', 'herbs': ['三棱', '莪术', '水蛭', '斑蝥', '穿山甲']},
        ]
    },
    {
        'id': 13, 'name': '化痰止咳平喘药',
        'subcategories': [
            {'name': '温化寒痰药', 'herbs': ['天南星', '半夏', '白附子', '白前', '旋覆花', '芥子', '猫爪草', '皂荚']},
            {'name': '清化热痰药', 'herbs': ['川贝母', '前胡', '浙贝母', '桔梗', '竹茹', '瓜蒌', '胖大海', '天竺黄', '昆布', '海藻', '瓦楞子', '竹沥', '黄药子', '海蛤壳', '海浮石', '礞石']},
            {'name': '止咳平喘药', 'herbs': ['百部', '紫菀', '桑白皮', '枇杷叶', '款冬花', '白果', '苦杏仁', '紫苏子', '葶苈子', '矮地茶', '马兜铃', '洋金花']},
        ]
    },
    {
        'id': 14, 'name': '安神药',
        'subcategories': [
            {'name': '重镇安神药', 'herbs': ['朱砂', '磁石', '龙骨', '琥珀']},
            {'name': '养心安神药', 'herbs': ['远志', '柏子仁', '酸枣仁', '灵芝', '合欢皮', '首乌藤']},
        ]
    },
    {
        'id': 15, 'name': '平肝息风药',
        'subcategories': [
            {'name': '平抑肝阳药', 'herbs': ['罗布麻叶', '石决明', '牡蛎', '珍珠母', '紫贝齿', '刺蒺藜', '代赭石']},
            {'name': '息风止痉药', 'herbs': ['天麻', '钩藤', '牛黄', '全蝎', '地龙', '羚羊角', '蜈蚣', '珍珠', '僵蚕']},
        ]
    },
    {
        'id': 16, 'name': '开窍药',
        'subcategories': [
            {'name': '开窍药', 'herbs': ['石菖蒲', '冰片', '麝香', '苏合香']},
        ]
    },
    {
        'id': 17, 'name': '补虚药',
        'subcategories': [
            {'name': '补气药', 'herbs': ['甘草', '人参', '山药', '太子参', '白术', '西洋参', '刺五加', '红景天', '黄芪', '党参', '大枣', '白扁豆', '绞股蓝', '蜂蜜', '饴糖', '沙棘']},
            {'name': '补阳药', 'herbs': ['巴戟天', '仙茅', '续断', '杜仲', '紫石英', '胡芦巴', '菟丝子', '补骨脂', '冬虫夏草', '海马', '鹿茸', '蛤蚧', '肉苁蓉', '淫羊藿', '锁阳', '阳起石', '海狗肾', '紫河车', '沙苑子', '蛤蟆油', '益智仁', '核桃仁', '韭菜子']},
            {'name': '补血药', 'herbs': ['白芍', '熟地黄', '何首乌', '当归', '龙眼肉', '阿胶']},
            {'name': '补阴药', 'herbs': ['北沙参', '玉竹', '百合', '黄精', '女贞子', '枸杞子', '龟甲', '鳖甲', '石斛', '黑芝麻', '麦冬', '天冬', '墨旱莲', '桑椹']},
        ]
    },
    {
        'id': 18, 'name': '收涩药',
        'subcategories': [
            {'name': '固表止汗药', 'herbs': ['麻黄根', '浮小麦', '糯稻根']},
            {'name': '敛肺涩肠药', 'herbs': ['五味子', '肉豆蔻', '诃子', '五倍子', '罂粟壳', '禹余粮', '赤石脂', '石榴皮', '乌梅']},
            {'name': '固精缩尿止带药', 'herbs': ['椿皮', '鸡冠花', '山茱萸', '芡实', '金樱子', '莲子', '覆盆子', '桑螵蛸', '海螵蛸', '刺猬皮']},
        ]
    },
    {
        'id': 20, 'name': '涌吐药',
        'subcategories': [
            {'name': '涌吐药', 'herbs': ['常山', '藜芦', '胆矾', '瓜蒂']},
        ]
    },
    {
        'id': 21, 'name': '攻毒杀虫止痒药',
        'subcategories': [
            {'name': '攻毒杀虫止痒药', 'herbs': ['土荆皮', '雄黄', '蛇床子', '蟾酥', '硫黄', '樟脑', '砒石', '白矾']},
        ]
    },
    {
        'id': 22, 'name': '拔毒化腐生肌药',
        'subcategories': [
            {'name': '拔毒化腐生肌药', 'herbs': ['炉甘石', '铅丹', '轻粉', '硼砂', '红粉']},
        ]
    },
]


def build_herb_url_map():
    """
    从中药主页提取所有药物的URL映射。

    主页结构：每个药物是一个 <a href="/traditionaldetails?id=XXX">药名</a> 链接
    返回：{'白芷': 'https://...id=41', '防风': 'https://...id=61', ...}
    """
    resp = requests.get(f'{BASE_URL}/traditional', timeout=15, headers=HEADERS)
    resp.encoding = 'utf-8'
    soup = BeautifulSoup(resp.text, 'lxml')

    herb_map = {}
    for a in soup.find_all('a'):
        href = a.get('href', '')
        text = a.get_text(strip=True)
        if 'traditionaldetails' in href and text:
            url = BASE_URL + href if not href.startswith('http') else href
            herb_map[text] = url

    return herb_map


def fetch_herb_detail(url):
    """抓取单个药物详情页"""
    resp = requests.get(url, timeout=15, headers=HEADERS)
    resp.encoding = 'utf-8'
    soup = BeautifulSoup(resp.text, 'lxml')
    text = soup.get_text(separator='\n', strip=True)
    return text


def parse_herb_text(text):
    """
    解析药物详情页的纯文本，提取结构化数据。

    页面结构（按顺序）：
      药品名称 → 始载于 → 别名 → 性味归经 → 功效 → 药材简介
      → 用法用量 → 注意事项 → 应用（编号列表）→ 相关配伍 → 相关方剂 → 炮制

    返回 dict，包含：
      基本信息：name, pinyin, source, aliases, nature_taste_channel, effects, introduction, usage_dosage, cautions
      扩展信息：applications[], pairings[], prescriptions[], processing_modern, processing_ancient
    """
    lines = text.split('\n')

    herb = {
        'name': '', 'pinyin': '', 'source': '', 'aliases': '',
        'nature_taste_channel': '', 'effects': '', 'introduction': '',
        'usage_dosage': '', 'cautions': '',
        'applications': [], 'pairings': [], 'prescriptions': [],
        'processing_modern': '', 'processing_ancient': '',
    }

    # ---- 提取基本信息字段 ----
    # 页面格式：字段名 → "：" → 字段值（间隔1行）
    for i, line in enumerate(lines):
        if line == '药品名称' and i + 2 < len(lines):
            val = lines[i + 2]
            herb['name'] = val.split('[')[0].strip() if '[' in val else val
            herb['pinyin'] = val
        elif line == '始载于' and i + 2 < len(lines):
            herb['source'] = lines[i + 2]
        elif line == '别名' and i + 2 < len(lines):
            herb['aliases'] = lines[i + 2]
        elif line == '性味归经' and i + 2 < len(lines):
            herb['nature_taste_channel'] = lines[i + 2]
        elif line == '功效' and i + 2 < len(lines) and lines[i + 1] == '：':
            herb['effects'] = lines[i + 2]
        elif line == '药材简介' and i + 2 < len(lines):
            herb['introduction'] = lines[i + 2]
        elif line == '用法用量' and i + 2 < len(lines):
            usage = []
            j = i + 2
            while j < len(lines) and lines[j] not in ('注意事项', '应用', '炮制', '基本信息', '药品名称'):
                if lines[j].strip():
                    usage.append(lines[j])
                j += 1
            herb['usage_dosage'] = '\n'.join(usage)
        elif line == '注意事项' and i + 2 < len(lines):
            cautions = []
            j = i + 2
            while j < len(lines) and lines[j] not in ('应用', '炮制', '基本信息', '药品名称'):
                if lines[j].strip():
                    cautions.append(lines[j])
                j += 1
            herb['cautions'] = '\n'.join(cautions)

    # ---- 提取临床应用 ----
    # 格式：编号列表 "1. 治外感风寒..." "2. 治阳明头痛..."
    in_app = False
    for line in lines:
        if line == '应用':
            in_app = True
            continue
        if line in ('相关配伍', '炮制'):
            in_app = False
            continue
        if in_app and re.match(r'^\d+\.', line):
            herb['applications'].append(line)

    # ---- 提取相关配伍 ----
    # 格式：配伍标题（如"白芷配细辛"）+ 配伍说明文本
    in_pair = False
    current_pair = None
    for line in lines:
        if line == '相关配伍':
            in_pair = True
            continue
        if line == '炮制':
            in_pair = False
            current_pair = None
            continue
        if in_pair:
            if '配' in line and len(line) < 30 and '汤' not in line and '散' not in line and '丸' not in line:
                current_pair = {'title': line, 'content': ''}
                herb['pairings'].append(current_pair)
            elif current_pair and line not in ('相关配伍', '炮制'):
                is_rx = ('汤' in line or '散' in line or '丸' in line or '饮' in line) and len(line) < 20
                if is_rx:
                    current_pair = None
                elif current_pair:
                    if current_pair['content']:
                        current_pair['content'] += '\n' + line
                    else:
                        current_pair['content'] = line

    # ---- 提取相关方剂 ----
    # 格式：方剂名（如"九味羌活汤"）→ 药物组成 → 功能与主治
    seen_rx = set()  # 去重用
    for i, line in enumerate(lines):
        if ('汤' in line or '散' in line or '丸' in line or '饮' in line) and \
           '药物组成' not in line and '功能与主治' not in line and \
           len(line) < 20 and line not in seen_rx:
            j = i + 1
            comp = ''
            func = ''
            while j < len(lines) and j < i + 5:
                if '药物组成' in lines[j]:
                    comp = lines[j].split(':', 1)[-1].split('：', 1)[-1].strip()
                elif '功能与主治' in lines[j]:
                    func_parts = [lines[j].split(':', 1)[-1].split('：', 1)[-1].strip()]
                    k = j + 1
                    while k < len(lines):
                        nl = lines[k].strip()
                        if not nl or '药物组成' in nl or '功能与主治' in nl:
                            break
                        if nl in ('相关配伍', '炮制', '现代炮制', '古法炮制'):
                            break
                        if ('汤' in nl or '散' in nl or '丸' in nl or '饮' in nl) and len(nl) < 20:
                            break
                        func_parts.append(nl)
                        k += 1
                    func = ''.join(func_parts)
                    break
                j += 1
            if comp or func:
                herb['prescriptions'].append({'name': line, 'composition': comp, 'function': func})
                seen_rx.add(line)

    # ---- 提取炮制方法 ----
    # 分为"现代炮制"和"古法炮制"两部分
    if '炮制' in text:
        proc_section = text.split('炮制')[1]
        if '扫码下载' in proc_section:
            proc_section = proc_section.split('扫码下载')[0]
        if '现代炮制' in proc_section:
            parts = proc_section.split('现代炮制')
            modern = parts[1] if len(parts) > 1 else ''
            if '古法炮制' in modern:
                modern = modern.split('古法炮制')[0]
            herb['processing_modern'] = modern.strip().lstrip('：:').strip()
        if '古法炮制' in proc_section:
            ancient = proc_section.split('古法炮制')[1]
            herb['processing_ancient'] = ancient.strip().lstrip('：:').strip()

    return herb


def load_progress():
    """加载抓取进度（断点续传用）。格式：{药名: {status: 'ok', data: {...}}}"""
    if os.path.exists(PROGRESS_JSON):
        with open(PROGRESS_JSON, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def save_progress(progress):
    """保存抓取进度"""
    with open(PROGRESS_JSON, 'w', encoding='utf-8') as f:
        json.dump(progress, f, ensure_ascii=False)


# ============================================================
#  第二部分：Word文档生成
# ============================================================

def set_cell_shading(cell, color):
    """设置Word表格单元格的背景色（用于基本信息表格的标签列）"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading_elm.append(shading)


def add_heading(doc, text, level=1):
    """添加带统一深绿色的标题（H1=大分类, H2=子分类, H4=药物名）"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)
    return h


def add_section_title(doc, text):
    """添加带底部蓝色下划线的小节标题（如"基本信息"、"临床应用"）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '0d47a1'})
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_herb_to_doc(doc, herb, herb_idx):
    """
    将单个药物的完整信息添加到文档。

    内容包括：
      - H4标题：序号 + 药名
      - 基本信息表格（药品名称、始载于、别名、性味归经、功效、药材简介、用法用量、注意事项）
      - 临床应用（编号列表）
      - 相关配伍（配对名称 + 说明）
      - 相关方剂（方名 + 药物组成 + 功能与主治）
      - 炮制（现代炮制 + 古法炮制）
    """
    add_heading(doc, f'{herb_idx}. {herb["name"]}', level=4)

    # Basic info table
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(13.5)

    fields = [
        ('药品名称', herb.get('pinyin', '')),
        ('始载于', herb.get('source', '')),
        ('别名', herb.get('aliases', '')),
        ('性味归经', herb.get('nature_taste_channel', '')),
        ('功效', herb.get('effects', '')),
        ('药材简介', herb.get('introduction', '')),
        ('用法用量', herb.get('usage_dosage', '')),
        ('注意事项', herb.get('cautions', '')),
    ]

    for label, value in fields:
        if value:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
            set_cell_shading(row.cells[0], 'e8f5e9')
            for p in row.cells[1].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # Applications
    if herb.get('applications'):
        add_section_title(doc, '临床应用')
        for app in herb['applications']:
            p = doc.add_paragraph(app, style='List Number')
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9.5)

    # Pairings
    if herb.get('pairings'):
        add_section_title(doc, '相关配伍')
        for pairing in herb['pairings']:
            p = doc.add_paragraph()
            r = p.add_run(pairing['title'])
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x2e, 0x7d, 0x32)
            if pairing.get('content'):
                p2 = doc.add_paragraph(pairing['content'])
                p2.paragraph_format.first_line_indent = Cm(0.75)
                p2.paragraph_format.space_after = Pt(4)
                for r2 in p2.runs:
                    r2.font.size = Pt(9.5)

    # Prescriptions
    if herb.get('prescriptions'):
        add_section_title(doc, '相关方剂')
        for rx in herb['prescriptions']:
            p = doc.add_paragraph()
            r = p.add_run(rx['name'])
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xbf, 0x36, 0x0c)
            if rx.get('composition'):
                p2 = doc.add_paragraph()
                r1 = p2.add_run('药物组成：')
                r1.bold = True
                r1.font.size = Pt(9)
                r2 = p2.add_run(rx['composition'])
                r2.font.size = Pt(9)
            if rx.get('function'):
                p3 = doc.add_paragraph()
                r1 = p3.add_run('功能与主治：')
                r1.bold = True
                r1.font.size = Pt(9)
                r2 = p3.add_run(rx['function'])
                r2.font.size = Pt(9)
                p3.paragraph_format.space_after = Pt(6)

    # Processing
    if herb.get('processing_modern') or herb.get('processing_ancient'):
        add_section_title(doc, '炮制')
        if herb.get('processing_modern'):
            p = doc.add_paragraph()
            r1 = p.add_run('现代炮制：')
            r1.bold = True
            r1.font.size = Pt(9.5)
            r2 = p.add_run(herb['processing_modern'])
            r2.font.size = Pt(9)
        if herb.get('processing_ancient'):
            p = doc.add_paragraph()
            r1 = p.add_run('古法炮制：')
            r1.bold = True
            r1.font.size = Pt(9.5)
            r2 = p.add_run(herb['processing_ancient'])
            r2.font.size = Pt(9)


def generate_docx(all_herb_data):
    """
    生成完整的Word文档。

    文档结构（3级导航）：
      H1  一、解表药
        H2  发散风寒药
          H4  1. 白芷
          H4  2. 防风
          ...
        H2  发散风热药
          H4  18. 升麻
          ...
      H1  二、清热药
        ...

    参数：
      all_herb_data — 所有药物的数据列表（来自 parse_herb_text）
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading('知源中医 - 常用中药大全', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('数据来源：www.zhiyuanzhongyi.com')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.italic = True

    doc.add_paragraph()

    # Build herb lookup
    herb_lookup = {h['name']: h for h in all_herb_data}

    # Chinese numerals for categories
    cn_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
               '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十', '二十一', '二十二']

    global_herb_idx = 0

    for cat_idx, cat in enumerate(CATEGORIES):
        # H1: Category
        add_heading(doc, f'{cn_nums[cat_idx]}、{cat["name"]}', level=1)

        for subcat in cat['subcategories']:
            # H2: Subcategory
            add_heading(doc, subcat['name'], level=2)

            # Fetch category page for description (only once per category)
            # We'll skip category descriptions to keep the doc focused on herbs

            for herb_name in subcat['herbs']:
                global_herb_idx += 1
                herb_data = herb_lookup.get(herb_name)
                if herb_data:
                    add_herb_to_doc(doc, herb_data, global_herb_idx)
                else:
                    p = doc.add_paragraph(f'{global_herb_idx}. {herb_name}（未抓取到数据）')
                    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                doc.add_paragraph()  # spacing between herbs

        if cat_idx < len(CATEGORIES) - 1:
            doc.add_page_break()

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'共收录 {global_herb_idx} 味中药')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT_DOCX)
    print(f'\n文档已生成: {OUTPUT_DOCX}')
    print(f'文件大小: {os.path.getsize(OUTPUT_DOCX) / 1024 / 1024:.1f} MB')


def main():
    """
    主函数：抓取全部中药 → 生成Word文档。

    流程：
      1. 访问主页，获取所有药物的URL映射
      2. 遍历分类结构，逐个抓取每味药的详情页
      3. 按分类层级生成Word文档

    断点续传：
      进度保存在 %TEMP%/herbs_progress.json
      重新运行时会自动跳过已成功抓取的药物
    """
    print("=" * 60)
    print("知源中医 - 常用中药大全 全量抓取")
    print("=" * 60)

    # 1. 从主页获取所有药物的URL
    print("\n1. 获取药物URL映射...")
    herb_url_map = build_herb_url_map()
    print(f"   找到 {len(herb_url_map)} 个药物链接")

    # 统计需要抓取的总数
    total_herbs = sum(len(sc['herbs']) for cat in CATEGORIES for sc in cat['subcategories'])
    print(f"   需要抓取: {total_herbs} 味中药")

    # 加载已有进度（断点续传）
    progress = load_progress()
    done = sum(1 for v in progress.values() if v.get('status') == 'ok')
    print(f"   已完成: {done}")

    # 2. 逐个抓取药物详情（跳过已成功的）
    print(f"\n2. 开始抓取（每药间隔{REQUEST_DELAY}秒）")
    print("-" * 60)

    start_time = time.time()
    errors = 0

    # 三层循环：大分类 → 子分类 → 具体药物，逐个抓取
    for cat in CATEGORIES:
        for subcat in cat['subcategories']:
            for herb_name in subcat['herbs']:
                # 跳过已成功抓取的药物（断点续传核心逻辑）
                if herb_name in progress and progress[herb_name].get('status') == 'ok':
                    continue

                # 从URL映射表中查找该药物的详情页地址
                url = herb_url_map.get(herb_name)
                if not url:
                    print(f"   ✗ {herb_name} - 未找到URL")
                    progress[herb_name] = {'status': 'no_url'}
                    continue

                try:
                    # 请求详情页 → 解析文本 → 保存数据
                    text = fetch_herb_detail(url)
                    herb_data = parse_herb_text(text)
                    herb_data['name'] = herb_name
                    progress[herb_name] = {'status': 'ok', 'data': herb_data}
                    print(f"   ✓ {herb_name} | {herb_data['effects'][:40]}...")
                except Exception as e:
                    progress[herb_name] = {'status': 'error'}
                    print(f"   ✗ {herb_name} - {type(e).__name__}")
                    errors += 1

                # 每抓一味药立即保存进度，防止中断丢失数据
                save_progress(progress)
                # 礼貌延迟，避免请求过快被封
                time.sleep(REQUEST_DELAY)

    elapsed = time.time() - start_time
    ok_count = sum(1 for v in progress.values() if v.get('status') == 'ok')
    print(f"\n抓取完成: {ok_count}/{total_herbs} | 失败: {errors} | 耗时: {elapsed/60:.1f}分钟")

    # 3. 汇总所有成功抓取的药物数据（保持分类顺序）
    all_herb_data = []
    for cat in CATEGORIES:
        for subcat in cat['subcategories']:
            for herb_name in subcat['herbs']:
                entry = progress.get(herb_name, {})
                if entry.get('status') == 'ok' and entry.get('data'):
                    all_herb_data.append(entry['data'])

    # 4. 按分类层级生成Word文档（H1→H2→H3→H4）
    print(f"\n3. 生成Word文档（{len(all_herb_data)} 味中药）...")
    generate_docx(all_herb_data)

    print("\n完成！")


if __name__ == '__main__':
    main()
