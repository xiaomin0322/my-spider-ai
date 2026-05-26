"""
星恒题库 - 中医师承确有专长 导出Word文档

按照网站页面的分类结构（科目→子章节→题目）输出Word文档。
每道题包含：题型、题干、选项、正确答案、解析。

输入：xingheng_questions.json
输出：星恒题库_中医师承确有专长.docx
"""

import json
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = r'C:\Users\zengmin.zhang\AI'
INPUT_JSON = os.path.join(OUTPUT_DIR, 'xingheng_questions.json')
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, '星恒题库_中医师承确有专长.docx')

# 答案颜色
ANSWER_COLOR = RGBColor(0x2e, 0x7d, 0x32)      # 绿色
EXPLANATION_COLOR = RGBColor(0x0d, 0x47, 0xa1)  # 蓝色
GRAY = RGBColor(0x99, 0x99, 0x99)
DARK = RGBColor(0x33, 0x33, 0x33)


def add_question(doc, q, idx):
    """添加一道题目"""
    # 题号 + 题型 + 题干
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    run_num = p.add_run(f'{idx}. ')
    run_num.bold = True
    run_num.font.size = Pt(11)

    if q.get('type'):
        run_type = p.add_run(f'【{q["type"]}】')
        run_type.font.size = Pt(10)
        run_type.font.color.rgb = GRAY
        p.add_run(' ')

    run_stem = p.add_run(q.get('stem', ''))
    run_stem.font.size = Pt(11)

    # 选项
    for opt in q.get('options', []):
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Cm(0.8)
        op.paragraph_format.space_after = Pt(1)

        run_label = op.add_run(f'{opt["label"]}. ')
        run_label.bold = True
        run_label.font.size = Pt(10.5)

        run_text = op.add_run(opt['text'])
        run_text.font.size = Pt(10.5)

    # 正确答案（绿色加粗）
    if q.get('answer'):
        pa = doc.add_paragraph()
        pa.paragraph_format.space_before = Pt(4)
        run_label = pa.add_run('【答案】')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = ANSWER_COLOR
        run_ans = pa.add_run(q['answer'])
        run_ans.bold = True
        run_ans.font.size = Pt(12)
        run_ans.font.color.rgb = ANSWER_COLOR

    # 解析（蓝色）
    if q.get('explanation'):
        pe = doc.add_paragraph()
        pe.paragraph_format.space_before = Pt(2)
        pe.paragraph_format.space_after = Pt(6)
        run_label = pe.add_run('【解析】')
        run_label.bold = True
        run_label.font.size = Pt(10.5)
        run_label.font.color.rgb = EXPLANATION_COLOR
        run_exp = pe.add_run(q['explanation'])
        run_exp.font.size = Pt(10)
        run_exp.font.color.rgb = DARK

    # 分隔线
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(2)
    sep.add_run('─' * 50).font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)


def main():
    with open(INPUT_JSON, 'r', encoding='utf-8') as f:
        all_questions = json.load(f)

    # 按科目→子章节分组
    structure = {}  # {subject: {chapter: [questions]}}
    for q in all_questions:
        subj = q.get('subject', '未知')
        chap = q.get('chapter', '未知')
        if subj not in structure:
            structure[subj] = {}
        if chap not in structure[subj]:
            structure[subj][chap] = []
        structure[subj][chap].append(q)

    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    # 文档标题
    title = doc.add_heading('星恒题库 · 中医师承确有专长', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)

    # 统计
    total = len(all_questions)
    has_answer = sum(1 for q in all_questions if q.get('answer'))
    total_chapters = sum(len(chaps) for chaps in structure.values())

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'共 {len(structure)} 个科目 | {total_chapters} 个章节 | {total} 道题')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'有答案 {has_answer} 题 | 数据来源：www.xinghengclass.com')
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY

    doc.add_paragraph()

    # 目录概览
    doc.add_heading('目录', level=1)
    for subj_idx, subj_name in enumerate(structure, 1):
        total_subj = sum(len(qs) for qs in structure[subj_name].values())
        p = doc.add_paragraph()
        run = p.add_run(f'{subj_idx}. {subj_name}')
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(f'  （{len(structure[subj_name])} 个章节，{total_subj} 题）')
        run2.font.size = Pt(10)
        run2.font.color.rgb = GRAY

        for chap_idx, chap_name in enumerate(structure[subj_name], 1):
            cp = doc.add_paragraph()
            cp.paragraph_format.left_indent = Cm(1)
            cnt = len(structure[subj_name][chap_name])
            run = cp.add_run(f'{subj_idx}.{chap_idx} {chap_name}（{cnt}题）')
            run.font.size = Pt(10)

    doc.add_page_break()

    # 逐科目输出
    for subj_idx, (subj_name, chapters) in enumerate(structure.items(), 1):
        total_subj = sum(len(qs) for qs in chapters.values())
        doc.add_heading(f'{subj_name}', level=1)

        info = doc.add_paragraph()
        info.add_run(f'共 {len(chapters)} 个章节，{total_subj} 道题').font.color.rgb = GRAY
        doc.add_paragraph()

        # 逐子章节输出
        for chap_idx, (chap_name, questions) in enumerate(chapters.items(), 1):
            doc.add_heading(f'{subj_idx}.{chap_idx} {chap_name}', level=2)

            chap_info = doc.add_paragraph()
            chap_info.add_run(f'共 {len(questions)} 题').font.color.rgb = GRAY

            for idx, q in enumerate(questions, 1):
                add_question(doc, q, idx)

        # 每个科目后分页
        if subj_idx < len(structure):
            doc.add_page_break()

    doc.save(OUTPUT_DOCX)
    size_kb = os.path.getsize(OUTPUT_DOCX) / 1024
    size_mb = size_kb / 1024
    print(f'文档已生成: {OUTPUT_DOCX}')
    if size_mb > 1:
        print(f'文件大小: {size_mb:.1f} MB')
    else:
        print(f'文件大小: {size_kb:.1f} KB')
    print(f'总题数: {total}')


if __name__ == '__main__':
    main()
